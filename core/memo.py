#!/usr/bin/env python
# encoding: utf-8
# Author: Johnny Loo
# Email: johnnyloo1985@gmail.com


import pickle
import os
from docx import Document
from .config import MemoConfig
from .log import MemoLog
from .memoTime import MemoTime
from .mail import MailSender
from .common import *


class Memo:
    """备忘录"""
    __slots__ = ('__id', 'date', 'thing')

    def __init__(self, id, date, thing):
        """定义备忘录属性"""
        self.__id = id
        self.date = date
        self.thing = thing

    @property
    def id(self):
        """id属性为只读"""
        return self.__id

    @id.setter
    def id(self, id):
        """设置id，用于删除一条备忘事项后的ID调整"""
        self.__id = id


class MemoAdmin:
    """备忘录功能"""

    def __init__(self, username):
        """初始化备忘录列表和菜单"""
        # 扩展：增加is_logout属性，默认为False，如果用户选择登出，则将其置为True
        self.is_logout = False
        # 扩展：传入username，根据它查找其备忘录数据文件路径，并读出其数据，供其他方法增删改查
        # 1. 创建主配置文件对象，查出其数据文件路径
        # 2. 读出其数据到memo_list列表中，如果没有数据文件，则创建之。

        mc = MemoConfig()

        self.memo_path = mc.get_memopath(username)
        self.user = username

        # 初始化日志logger
        self.ml = MemoLog('MemoModule')

        try:
            self.load()
        except:
            self.memo_list = []
            self.save()
        finally:
            self.__menu = {
                'D': '查看',
                'A': '添加',
                'R': '移除',
                'C': '修改',
                'E': '导出',
                'S': '分享',
                'Q': '登出'
            }

    def deal_id(self, id_str):
        """处理用户输入的ID"""
        try:
            id = int(id_str)
            if 0 < id <= len(self.memo_list):
                self.ml.logger.debug('The id inputted is legal.')
                return id
            else:
                raise ValueError
        except Exception as e:
            self.ml.logger.debug('The id inputted is illegal, ', e)
            return 'error'

    def deal_entry(self, entry):
        """处理用户输入的事项"""
        try:
            if not entry:
                raise ValueError('输入为空')
            date = MemoTime.extractTime(entry)  # 从字符串中提取日期返回标准格式
            self.ml.logger.debug('Extracting date from the entry is successful.')
            return date, entry
        except ValueError:
            self.ml.logger.debug('The entry inputted is empty.')
            return 'empty'

    def display(self):
        """打印备忘录列表"""
        try:
            print('-' * 65)
            print(Color.blue('ID'.center(5)), Color.blue('时间'.center(30)), Color.blue('事件'.center(30)))
            print('-' * 65)
            self.load()
            if not self.memo_list:
                print(Color.yellow('当前无事项！'))
                self.ml.logger.debug(f'User {self.user} has no entries.')
            else:
                for m in self.memo_list:
                    print(Color.purple(str(m.id).center(5)), Color.purple(m.date.center(30)), Color.purple(m.thing.ljust(30)))
                print('-' * 65)
                self.ml.logger.debug(f'User {self.user} is printing entries.')
                return apiReturn()
        except Exception as e:
            self.ml.logger.debug(f'User {self.user} printing entries failed.')
            return apiReturn(1, FAILURE)

    @classmethod
    def query(cls, user, year, month):
        """根据参数返回指定时段的备忘录事项"""
        try:
            ma = cls(user)
        except:
            return apiReturn(2, FAILURE)

        try:
            ma.load()
            patternTime = re.compile(r'(\d{4})年(\d{1,2})月(\d{1,2})日.+')
            data = []
            for memo in ma.memo_list:
                if patternTime.match(memo.date):
                    y = patternTime.match(memo.date).group(1)
                    m = patternTime.match(memo.date).group(2)
                    if y == year and m == month.rjust(2,'0'):
                        entry = {}
                        entry['id'] = memo.id
                        entry['date'] = memo.date
                        entry['thing'] = memo.thing
                        data.append(entry)
            return apiReturn(data=data)
        except Exception as e:
            return apiReturn(1,FAILURE)

    def add(self):
        """添加一条备忘事项"""
        while True:
            entry_list = self.deal_entry(input('请输入备忘事项：').strip())
            if entry_list == 'empty':
                print(Color.yellow('输入为空，请重新输入。'))
            else:
                break
        try:
            date, thing = entry_list
            id = len(self.memo_list) + 1
            self.memo_list.append(Memo(id, date, thing))
            self.save()
            self.ml.logger.info(f'User {self.user} is adding entry successfully')
            return apiReturn()
        except Exception as e:
            self.ml.logger.error(f'User {self.user} is adding entry failed, ', e)
            return apiReturn(1,FAILURE)


    def remove(self):
        """移除一条备忘事项"""
        self.display()
        while True:
            id = self.deal_id(input('请输入要移除事项的ID：').strip())
            if id == 'error':
                print(Color.yellow('非法输入，请重新输入。'))
            else:
                break
        index = id - 1
        try:
            del self.memo_list[index]
            for memo in self.memo_list:
                if memo.id > id:
                    memo.id -= 1
            self.save()
            self.ml.logger.info(f'User {self.user} removed an entry successfully.')
            return apiReturn()
        except Exception as e:
            self.ml.logger.error(f'User {self.user} removed an entry failed, ', e)
            return apiReturn(1, FAILURE)




    def change(self):
        """修改一条备忘事项"""
        self.display()
        while True:
            id = self.deal_id(input('请输入要修改事项的ID：').strip())
            if id == 'error':
                print(Color.yellow('非法输入，请重新输入。'))
            else:
                break
        index = id - 1
        while True:
            entry_list = self.deal_entry(input('请输入备忘事项：').strip())
            if entry_list == 'empty':
                print(Color.yellow('输入为空，请重新输入。'))
            else:
                break
        try:
            date, thing = entry_list
            memo_new = Memo(id, date, thing)
            self.memo_list[index] = memo_new
            self.save()
            self.ml.logger.info(f'User {self.user} changed an entry successfully.')
            return apiReturn()
        except Exception as e:
            self.ml.logger.error(f'User {self.user} changed an entry failed, ', e)
            return apiReturn(1, FAILURE)


    def save(self):
        """保存备忘录列表"""
        with open(self.memo_path, 'wb') as f:
            pickle.dump(self.memo_list, f)

    @classmethod
    def mail(cls, user, year, month):
        """发送用户的备忘录信息到他的邮箱中"""
        # - 加载备忘录信息
        # - 加载用户邮件相关配置
        # - 发邮件
        try:
            mc = MemoConfig()
            if not mc.cfg.has_section(user):
                raise ValueError(f'{user}不存在')
        except ValueError as ve:
            return apiReturn(2,FAILURE)

        try:
            mailCfg = mc.getMailcfg(user)
            basedir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
            senderAuthPath = os.path.join(basedir, 'cfg', mailCfg['senderAuth'])
            with open(senderAuthPath, 'rb') as f:
                password = pickle.load(f)
            ms = MailSender(mailCfg['smtpSvr'], mailCfg['sender'], password)
            ms.addRecevicer(mailCfg['recevicer'])
            # 获得并创建邮件正文备忘录信息
            bodyHeader = f'<h1>{user}的备忘录</h1>'
            bodyBefore = """
                        <table border="1">
                            <tr>
                                <th>ID</th>
                                <th>时间</th>
                                <th>事件</th>
                            </tr>
                        """
            bodyAfter = """
                        </table>
                        """
            ma = MemoAdmin(user)
            ma.load()
            if not ma.memo_list:
                bodyMiddle = """
                                <tr>
                                    <td>当前无事项！</td>
                                    <td>当前无事项！</td>
                                    <td>当前无事项！</td>
                                </tr>
                            """
            else:
                patternTime = re.compile(r'(\d{4})年(\d{1,2})月(\d{1,2})日.+')
                bodyMiddle = ''
                for memo in ma.memo_list:
                    if patternTime.match(memo.date):
                        y = patternTime.match(memo.date).group(1)
                        m = patternTime.match(memo.date).group(2)
                        if y == year and m == month.rjust(2, '0'):
                            bodyEntry = f"""
                                            <tr>
                                                <td>{str(memo.id)}</td>
                                                <td>{memo.date}</td>
                                                <td>{memo.thing}</td>
                                            </tr>
                                        """
                            bodyMiddle += bodyEntry

            body = bodyHeader + bodyBefore + bodyMiddle + bodyAfter

            ms.write('html', subject=f'来自{user}的备忘录分享', body=body)
            ms.send()
            return apiReturn()
        except:
            return apiReturn(1,FAILURE)


    def share(self):
        """分享备忘录"""
        try:
            recevicer = checkMail('您想分享给谁？请输入TA的Email:')
            mailCfg = MemoConfig().getMailcfg(self.user)
            basedir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
            senderAuthPath = os.path.join(basedir, 'cfg', mailCfg['senderAuth'])
            with open(senderAuthPath, 'rb') as f:
                password = pickle.load(f)
            ms = MailSender(mailCfg['smtpSvr'], mailCfg['sender'], password)
            ms.addRecevicer(recevicer)
            # 获得并创建邮件正文备忘录信息
            bodyHeader = f'<h1>{self.user}的备忘录</h1>'
            bodyBefore = """
            <table border="1">
                <tr>
                    <th>ID</th>
                    <th>时间</th>
                    <th>事件</th>
                </tr>
            """
            bodyAfter = """
            </table>
            """
            self.load()
            if not self.memo_list:
                bodyMiddle = """
                    <tr>
                        <td>当前无事项！</td>
                        <td>当前无事项！</td>
                        <td>当前无事项！</td>
                    </tr>
                """
            else:
                bodyMiddle = ''
                for m in self.memo_list:
                    bodyEntry = f"""
                        <tr>
                            <td>{str(m.id)}</td>
                            <td>{m.date}</td>
                            <td>{m.thing}</td>
                        </tr>
                    """
                    bodyMiddle += bodyEntry

            body = bodyHeader + bodyBefore + bodyMiddle + bodyAfter

            ms.write('html', subject=f'来自{self.user}的备忘录分享', body=body)
            ms.send()
            print(Color.green('分享成功！'))
        except Exception as e:
            print(Color.red('报错：'+e))



    def run(self):
        """主程序"""
        while True:
            print(Color.purple(f'当前用户：{self.user}'.center(30, '*')))
            for k, v in self.__menu.items():
                print(Color.green(f'{k}：{v}'.ljust(10)), end='')
            print()
            choose = input('请选择：').strip().lower()
            if choose == 'd':
                self.display()
            elif choose == 'a':
                self.add()
            elif choose == 'r':
                self.remove()
            elif choose == 'c':
                self.change()
            elif choose == 'e':
                self.export()  # 导出备忘录列表
            elif choose == 's':
                self.share()  # 备忘录分享
            elif choose == 'q':
                self.is_logout = True
                break
            else:
                print(Color.red('非法输入，请重新输入。'))
                self.ml.logger.debug(f'User {self.user}\'s input is illegal.')
            self.ml.logger.debug(f'User {self.user}\'s input is legal.')


    def load(self):
        """读取备忘录列表"""
        with open(self.memo_path, 'rb') as f:
            self.memo_list = pickle.load(f)


    def export(self):
        """导出备忘录列表为docx文件"""
        """
        - 加载用户备忘录列表
        - 创建word文档对象
        - 添加title
        - 创建表格，写入备忘录信息
        - 保存word文档
        """
        try:
            doc = Document()
            doc.add_heading('备忘录列表',level=0)
            doc.add_paragraph(f'用户名：{self.user}', style='Subtitle')
            tab = doc.add_table(rows=len(self.memo_list)+1, cols=3, style='Medium List 1')
            row0 = tab.row_cells(0)
            row0[0].text = 'ID'
            row0[1].text = '日期'
            row0[2].text = '事件'
            if self.memo_list:
                for m in self.memo_list:
                    tab.cell(m.id, 0).text = str(m.id)
                    tab.cell(m.id, 1).text = m.date
                    tab.cell(m.id, 2).text = m.thing
            else:
                tab.add_row()
                c10 = tab.cell(1,0)
                c11 = tab.cell(1,1)
                c12 = tab.cell(1,2)
                c10.merge(c11)
                c10.merge(c12)
                c10.text = '没有事项'
            basedir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
            expdir = os.path.join(basedir, 'export')
            expfile = os.path.join(expdir, f'{self.user}_memo.docx')
            if not os.path.isdir(expdir):
                os.mkdir(expdir)
            doc.save(expfile)
            print(Color.blink(f'导出成功，文件路径：{expfile}'))
            self.ml.logger.info(f'User "{self.user}" exported memo successfully.')
            return apiReturn()
        except Exception as e:
            print('报错：', e)
            self.ml.logger.error(f'User "{self.user}" exported memo fail, ', e)
            return apiReturn(1, FAILURE)


def main():
    pass


if __name__ == '__main__':
    main()